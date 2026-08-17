"""ReportService — professional .docx generation.

Builds the full Patent Chemistry Analysis report:
  Cover page, Table of Contents, 12 numbered sections + appendix,
  chemical structure images embedded throughout, and the static flowchart
  auto-split across pages for print readability.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.config import settings
from backend.models import Compound, Job, Reaction, Stage
from backend.services.flowchart_service import build_graph, render_static

logger = logging.getLogger("pca.report")

MAROON = RGBColor(0x80, 0x00, 0x00)
DARK = RGBColor(0x3B, 0x2A, 0x20)
CREAM_BG = "FFF8E7"
GRAY = RGBColor(0x6B, 0x5B, 0x50)
NS = "Not specified in patent"


def generate_report(job: Job, out_dir: Path) -> Path:
    """Generate the Word report for a completed job. Returns the .docx path."""
    out_dir.mkdir(parents=True, exist_ok=True)
    safe = "".join(c if c.isalnum() or c in " .-_" else "_" for c in job.original_filename)
    out_path = out_dir / f"{safe}-analysis.docx"

    doc = Document()
    _style_document(doc)

    # ------------------------------------------------------------------ #
    # Cover page
    # ------------------------------------------------------------------ #
    _cover_page(doc, job)

    # ------------------------------------------------------------------ #
    # Table of contents (field-based so Word builds it on open)
    # ------------------------------------------------------------------ #
    _heading(doc, "Table of Contents", level=0)
    _insert_toc(doc)
    doc.add_page_break()

    # 1. Patent Overview
    _heading(doc, "1. Patent Overview", level=1)
    _kv(doc, "Patent title", job.patent_title)
    _kv(doc, "Patent number", job.patent_number)
    _kv(doc, "Assignee", job.assignee)
    _kv(doc, "Inventors", job.inventors)
    _kv(doc, "Filing date", job.filing_date)
    _kv(doc, "Analysis date", datetime.now().strftime("%Y-%m-%d %H:%M"))
    _kv(doc, "File analyzed", job.original_filename)
    _kv(doc, "Page count", str(job.page_count))

    # 2. Executive Summary
    _heading(doc, "2. Executive Summary", level=1)
    qc = (job.extra or {}).get("qc_issues", [])
    score = (job.extra or {}).get("qc_score", 1.0)
    _p(doc, (
        f"This report analyzes the chemistry disclosed in \"{job.patent_title}\" "
        f"({job.patent_number}). The analysis identified {len(job.compounds)} chemical "
        f"entities and {len(job.reactions)} reactions across {job.page_count} pages. "
        f"Overall quality score: {score:.0%}. "
        + ("Database conflicts or unresolved structures are flagged in the relevant sections."
           if qc else "All key structures were validated against public chemistry databases.")
    ))

    # 3. Chemical Entities
    _heading(doc, "3. Chemical Entities", level=1)
    _compounds_table(doc, job.compounds)

    # 4. Complete Reaction Flowchart
    _heading(doc, "4. Complete Reaction Flowchart", level=1)
    _p(doc, "The complete reaction pathway is shown below. The flowchart is split "
            "into sections if it is larger than a single printed page.")
    _flowchart_pages(doc, job)

    # 5. Reaction-by-Reaction Analysis
    _heading(doc, "5. Reaction-by-Reaction Analysis", level=1)
    for r in sorted(job.reactions, key=lambda x: x.rid):
        _reaction_section(doc, r, job)

    # 6. Stage-by-Stage Explanation
    _heading(doc, "6. Stage-by-Stage Explanation", level=1)
    for s in sorted(job.stages, key=lambda x: x.order_idx):
        _stage_section(doc, s)

    # 7. Chemical Transformation Analysis
    _heading(doc, "7. Chemical Transformation Analysis", level=1)
    for r in sorted(job.reactions, key=lambda x: x.rid):
        _transform_section(doc, r)

    # 8. Manufacturing Process
    _heading(doc, "8. Manufacturing Process", level=1)
    _manufacturing_section(doc, job)

    # 9. Equipment and Machinery
    _heading(doc, "9. Equipment and Machinery", level=1)
    _equipment_section(doc, job)

    # 10. Manufacturer Details
    _heading(doc, "10. Manufacturer Details", level=1)
    _kv(doc, "Manufacturer / Assignee", job.assignee)
    _kv(doc, "Applicant", job.assignee)
    _kv(doc, "Inventors", job.inventors)
    _kv(doc, "Organization", job.assignee)
    _kv(doc, "Country", "Not specified in patent")
    _kv(doc, "Manufacturing facility", "Not specified in patent")

    # 11. Chemistry Database References
    _heading(doc, "11. Chemistry Database References", level=1)
    _db_refs(doc, job)

    # 12. Patent References
    _heading(doc, "12. Patent References", level=1)
    _kv(doc, "Patent", f"{job.patent_number} — {job.patent_title}")
    _kv(doc, "Page references", ", ".join(sorted({str(c.source_page) for c in job.compounds if c.source_page})) or NS)

    # Appendix
    _heading(doc, "Appendix", level=1)
    _p(doc, "Confidence indicators and data-source labels used in this report:")
    _bullet(doc, "\u201cPatent-stated\u201d — directly quoted from the patent text (page-referenced).")
    _bullet(doc, "\u201cChemistry database\u201d — verified via OPSIN / PubChem / NIH CIR.")
    _bullet(doc, "\u201cAI Chemical Interpretation\u201d — inferred reasoning, clearly labelled.")
    _bullet(doc, "\u201cNot specified in patent\u201d — the patent does not disclose this information.")

    doc.save(str(out_path))
    logger.info("report written: %s", out_path)
    return out_path


# --------------------------------------------------------------------------- #
# Section builders
# --------------------------------------------------------------------------- #
def _cover_page(doc: Document, job: Job) -> None:
    for _ in range(4):
        doc.add_paragraph()
    _p(doc, "Patent Chemistry Analysis", size=30, bold=True, color=MAROON,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, job.patent_title, size=16, bold=False, color=DARK,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, job.patent_number, size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, datetime.now().strftime("%B %d, %Y"), size=12, color=GRAY,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def _compounds_table(doc: Document, compounds: list[Compound]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(["ID", "Name", "Formula", "MW", "Role", "Page"]):
        hdr[i].text = t
    for c in compounds:
        row = table.add_row().cells
        row[0].text = c.cid
        row[1].text = c.name
        row[2].text = c.molecular_formula
        row[3].text = f"{c.molecular_weight:.1f}" if c.molecular_weight else NS
        row[4].text = c.role
        row[5].text = str(c.source_page) if c.source_page else NS
    doc.add_paragraph()


def _reaction_section(doc: Document, r: Reaction, job: Job) -> None:
    _heading(doc, f"{r.rid} — {r.name or r.type}", level=2)
    _kv(doc, "Reaction type", r.type)
    _kv(doc, "Reactants", r.reactants_text or NS)
    _kv(doc, "Products", r.products_text or NS)
    _kv(doc, "Reagents", r.reagents)
    _kv(doc, "Catalysts", r.catalysts)
    _kv(doc, "Solvents", r.solvents)
    _kv(doc, "Temperature", r.temperature)
    _kv(doc, "Pressure", r.pressure)
    _kv(doc, "Time", r.time)
    _kv(doc, "Atmosphere", r.atmosphere)
    _kv(doc, "Yield", r.yield_pct)
    _kv(doc, "Work-up", r.workup)
    _kv(doc, "Equipment", (r.extra or {}).get("equipment", NS))
    _kv(doc, "Patent evidence", f"Patent page {r.source_page}" if r.source_page else NS)
    if r.source_text:
        _quote(doc, r.source_text)
    for p in r.participants:
        if p.smiles and p.structure_valid:
            _structure_image(doc, job, p)
    _p(doc, f"Confidence: {r.confidence:.0%} — {confidence_label(r.confidence)}",
       italic=True, color=GRAY)


def _stage_section(doc: Document, s: Stage) -> None:
    _heading(doc, s.title, level=2)
    _kv(doc, "Purpose", s.purpose)
    _kv(doc, "Starting material", s.starting_material)
    _kv(doc, "Reagents", s.reagents)
    _kv(doc, "Conditions", s.conditions)
    _kv(doc, "Reaction", s.reaction)
    _kv(doc, "Product", s.product)
    _kv(doc, "What changed?", s.what_changed)
    _kv(doc, "Why this step is required", s.why_required)
    _kv(doc, "What happens chemically", s.chemistry)
    _kv(doc, "Purification", s.purification)
    _kv(doc, "Yield", s.yield_value)
    _kv(doc, "Equipment", s.equipment)
    _kv(doc, "Patent reference", s.patent_ref)
    _kv(doc, "Scale", s.scale)


def _transform_section(doc: Document, r: Reaction) -> None:
    import json

    wc = json.loads(r.what_changed) if r.what_changed else {}
    _heading(doc, f"{r.rid} — What Changed?", level=2)
    if not wc:
        _p(doc, "Not enough structural data to determine what changed.")
        return
    _bullet(doc, f"Reaction type: {', '.join(wc.get('reaction_types') or []) or 'Not specified in patent'}")
    if wc.get("formula_before") and wc.get("formula_after"):
        _bullet(doc, f"Molecular formula: {wc['formula_before']} \u2192 {wc['formula_after']}")
    if wc.get("atoms_added"):
        _bullet(doc, f"Atoms added: {', '.join(wc['atoms_added'])}")
    if wc.get("atoms_removed"):
        _bullet(doc, f"Atoms removed: {', '.join(wc['atoms_removed'])}")
    if wc.get("functional_group_added"):
        _bullet(doc, f"Functional groups introduced: {', '.join(wc['functional_group_added'])}")
    if wc.get("functional_group_removed"):
        _bullet(doc, f"Functional groups lost: {', '.join(wc['functional_group_removed'])}")
    for b in wc.get("bond_formed") or []:
        _bullet(doc, f"Bond formed: {b}")
    for b in wc.get("bond_broken") or []:
        _bullet(doc, f"Bond broken: {b}")
    if wc.get("similarity") is not None:
        _bullet(doc, f"Tanimoto similarity: {wc['similarity']:.2f}")
    _p(doc, f"(basis: {wc.get('basis', 'RDKit structural analysis')})", italic=True, color=GRAY)


def _manufacturing_section(doc: Document, job: Job) -> None:
    stages = sorted(job.stages, key=lambda x: x.order_idx)
    _kv(doc, "Scale summary", (job.extra or {}).get("scale_summary", NS))
    _kv(doc, "Laboratory-scale vs industrial", "Distinguished per stage in section 6.")
    if stages:
        _p(doc, "Process units and equipment are enumerated in section 9.")
    else:
        _p(doc, "No manufacturing process was reconstructed from this patent.")
    _kv(doc, "Raw materials / starting materials", "; ".join(
        {s.starting_material for s in stages if s.starting_material != NS} or [NS]))


def _equipment_section(doc: Document, job: Job) -> None:
    items: set[str] = set()
    for s in job.stages:
        if s.equipment != NS:
            items.update(x.strip() for x in s.equipment.replace(";", ",").split(",") if x.strip())
    for r in job.reactions:
        eq = (r.extra or {}).get("equipment")
        if eq and eq != NS:
            items.update(x.strip() for x in eq.replace(";", ",").split(",") if x.strip())
    if items:
        for it in sorted(items):
            _bullet(doc, it)
    else:
        _p(doc, "Not specified in patent")


def _db_refs(doc: Document, job: Job) -> None:
    providers = set()
    for c in job.compounds:
        for p in (c.extra or {}).get("providers", []):
            providers.add(p)
    _bullet(doc, "PubChem PUG REST — compound identity, formula, weight, InChI/InChIKey, CAS.")
    _bullet(doc, "OPSIN — IUPAC/chemical name to SMILES resolution.")
    _bullet(doc, "NIH CIR (Cactus) — identifier resolution fallback.")
    _bullet(doc, "RDKit 2026.03 — structure validation, descriptors, functional groups, MCS.")
    _bullet(doc, "ChEBI — ontology enrichment (biological roles).")
    _p(doc, "Providers actually consulted for this analysis: "
            + (", ".join(sorted(providers)) if providers else "none"), italic=True)


def _flowchart_pages(doc: Document, job: Job) -> None:
    nodes, edges = build_graph(job)
    if not nodes:
        _p(doc, "No flowchart could be generated (no resolved compounds/reactions).")
        return
    out_dir = settings.REPORT_DIR / f"flowchart_{job.id}"
    images = render_static(job, nodes, edges, out_dir)
    for img in images:
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.5))
            doc.add_page_break()


def _structure_image(doc: Document, job: Job, c: Compound) -> None:
    path = c.image_path
    if not path or not Path(path).exists():
        from backend.services.structure_service import render_structure
        from backend.agents.pipeline import job_img_dir

        p = job_img_dir(job.id) / f"{job.id}_{c.cid}.png"
        if render_structure(c.smiles or "", p):
            path = str(p)
    if path and Path(path).exists():
        doc.add_picture(str(path), width=Inches(2.2))
        caption = f"{c.cid} — {c.name} (Patent page {c.source_page})" if c.source_page else f"{c.cid} — {c.name}"
        _p(doc, caption, italic=True, color=GRAY)


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #
def _style_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width, section.page_height = Inches(8.5), Inches(11)
        section.orientation = WD_ORIENT.PORTRAIT
        for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, m, Inches(0.8))
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4)


def _heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = MAROON if level <= 1 else DARK
    run.font.size = Pt(22 if level == 0 else (15 if level == 1 else 12))
    p.paragraph_format.space_before = Pt(14 if level > 0 else 8)
    p.paragraph_format.space_after = Pt(6)


def _p(doc: Document, text: str, size: int = 10.5, bold: bool = False,
       color=GRAY, align=None, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    if align:
        p.alignment = align


def _kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    run_k = p.add_run(f"{key}: ")
    run_k.bold = True
    run_k.font.color.rgb = DARK
    run_v = p.add_run(str(value))
    run_v.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.color.rgb = DARK
    run.font.size = Pt(10.5)


def _quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text[:900])
    run.italic = True
    run.font.color.rgb = GRAY
    run.font.size = Pt(9.5)


def _insert_toc(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose 'Update Field' to populate."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r2 = para.add_run()
    r2._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)


def confidence_label(score: float) -> str:
    if score >= 0.9:
        return "High Confidence"
    if score >= 0.7:
        return "Moderate Confidence"
    if score >= 0.5:
        return "Low Confidence"
    return "Unverified"
